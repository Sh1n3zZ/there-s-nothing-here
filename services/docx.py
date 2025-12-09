from docx import Document
from typing import Dict
from io import BytesIO

def replace_keywords(file_content: bytes, replacements: Dict[str, str]) -> bytes:
    """
    Replace keywords in a docx document in memory.
    
    Args:
        file_content: Document content as bytes
        replacements: Dictionary mapping old text to new text
        
    Returns:
        Modified document content as bytes
    """
    document = Document(BytesIO(file_content))

    for paragraph in document.paragraphs:
        for old_text, new_text in replacements.items():
            if old_text in paragraph.text:

                for run in paragraph.runs:
                    if old_text in run.text:
                        run.text = run.text.replace(old_text, new_text)
    
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for old_text, new_text in replacements.items():
                        if old_text in paragraph.text:
                            for run in paragraph.runs:
                                if old_text in run.text:
                                    run.text = run.text.replace(old_text, new_text)

    output = BytesIO()
    document.save(output)
    output.seek(0)
    return output.getvalue()
